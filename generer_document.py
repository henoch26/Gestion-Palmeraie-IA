from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True, color=(31, 73, 125))
    # Trait sous le titre
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=(0, 112, 192))
    return p

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

def question_item(doc, num, question, detail=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"Q{num}. ")
    set_run_font(r1, size=11, bold=True, color=(192, 0, 0))
    r2 = p.add_run(question)
    set_run_font(r2, size=11, bold=True)
    if detail:
        pd = doc.add_paragraph()
        pd.paragraph_format.left_indent  = Cm(1.2)
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after  = Pt(3)
        rd = pd.add_run(detail)
        set_run_font(rd, size=10, italic=True, color=(89, 89, 89))

def feature_block(doc, icon, title, description, sub_items=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{icon}  {title}")
    set_run_font(r1, size=11, bold=True, color=(0, 112, 192))
    pd = doc.add_paragraph()
    pd.paragraph_format.left_indent  = Cm(1)
    pd.paragraph_format.space_before = Pt(0)
    pd.paragraph_format.space_after  = Pt(2)
    rd = pd.add_run(description)
    set_run_font(rd, size=10.5)
    if sub_items:
        for item in sub_items:
            pb = doc.add_paragraph(style="List Bullet")
            pb.paragraph_format.left_indent  = Cm(1.4)
            pb.paragraph_format.space_before = Pt(0)
            pb.paragraph_format.space_after  = Pt(1)
            rb = pb.add_run(item)
            set_run_font(rb, size=10)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 80)
    set_run_font(run, size=8, color=(189, 189, 189))

# ════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("UNIVERSITÉ FÉLIX HOUPHOUËT-BOIGNY")
set_run_font(r, size=13, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Master 2 – Bases de Données et Génie Logiciel")
set_run_font(r, size=11, italic=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("APPLICATION DE GESTION DE PALMERAIE")
set_run_font(r, size=20, bold=True, color=(0, 112, 192))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Projet de Mémoire — Thème 2")
set_run_font(r, size=13, italic=True)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Document de travail à destination du Directeur de mémoire")
set_run_font(r, size=11, color=(89, 89, 89))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Date : {datetime.date.today().strftime('%d/%m/%Y')}")
set_run_font(r, size=11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  PARTIE 1 — FONCTIONNALITÉS IMPLÉMENTÉES
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PARTIE 1")
set_run_font(r, size=10, bold=True, color=(127, 127, 127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FONCTIONNALITÉS DÉJÀ IMPLÉMENTÉES")
set_run_font(r, size=16, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Ce document présente l'ensemble des modules développés dans l'application "
    "de gestion de palmeraie. Son objectif est de donner une vue claire et complète "
    "de l'existant afin de faciliter les échanges avec le directeur de mémoire."
)
set_run_font(r, size=11, italic=True, color=(89, 89, 89))

doc.add_paragraph()

# ── 1. Authentification ──────────────────────────────────────────────────────
heading1(doc, "1. Module Authentification et Gestion des Comptes")

feature_block(
    doc, "🔐", "Connexion sécurisée",
    "Chaque utilisateur accède à l'application via un identifiant et un mot de passe. "
    "Un mécanisme de token JWT sécurise toutes les communications entre le navigateur et le serveur.",
    [
        "Page de connexion avec gestion des erreurs",
        "Déconnexion automatique après expiration du token",
        "Redirection automatique selon le rôle de l'utilisateur",
    ]
)

feature_block(
    doc, "🔑", "Changement de mot de passe obligatoire",
    "Lorsqu'un administrateur crée un nouveau compte, le système force l'utilisateur "
    "à changer son mot de passe dès sa première connexion.",
    [
        "Détection automatique du statut 'must_change_password'",
        "Page dédiée au changement de mot de passe",
    ]
)

feature_block(
    doc, "👤", "Profil utilisateur",
    "Chaque utilisateur dispose d'une page de profil où il peut consulter et modifier "
    "ses informations personnelles (nom, numéro de téléphone, mot de passe).",
)

# ── 2. Gestion des utilisateurs ──────────────────────────────────────────────
heading1(doc, "2. Module Gestion des Utilisateurs (Administration)")

body(
    doc,
    "Ce module est réservé aux administrateurs. Il permet de gérer tous les comptes "
    "utilisateurs de l'application."
)

feature_block(
    doc, "👥", "Deux rôles distincts",
    "L'application distingue deux profils d'utilisateurs avec des droits différents :",
    [
        "Administrateur : accès complet à toutes les fonctionnalités, gestion des comptes, validation des fiches.",
        "Superviseur : saisie et consultation des fiches de récolte et de travaux de son périmètre.",
    ]
)

feature_block(
    doc, "⚙️", "Administration des comptes",
    "L'administrateur peut créer, modifier, activer ou désactiver les comptes utilisateurs "
    "depuis une interface dédiée.",
    [
        "Création de nouveaux comptes superviseurs",
        "Attribution et modification des rôles",
        "Réinitialisation de mot de passe",
        "Activation / désactivation d'un compte",
        "Consultation de la liste complète des utilisateurs",
    ]
)

# ── 3. Gestion des secteurs ──────────────────────────────────────────────────
heading1(doc, "3. Module Gestion des Secteurs")

body(
    doc,
    "Un secteur représente une zone géographique délimitée de la palmeraie. "
    "Chaque secteur possède ses propres caractéristiques et est rattaché aux activités de récolte et de travaux."
)

feature_block(
    doc, "🗺️", "Fiche secteur complète",
    "Chaque secteur est enregistré avec l'ensemble de ses informations descriptives :",
    [
        "Code unique d'identification du secteur",
        "Nom et description",
        "Superficie en hectares",
        "Localisation / zone géographique",
        "Date de création dans le système",
    ]
)

feature_block(
    doc, "📋", "Liste et recherche des secteurs",
    "Une page liste tous les secteurs avec possibilité de recherche et de tri. "
    "Chaque ligne donne accès à la fiche détaillée du secteur."
)

feature_block(
    doc, "📊", "Détail et historique d'un secteur",
    "La page de détail d'un secteur affiche l'historique complet des récoltes "
    "et des travaux réalisés sur ce secteur, avec des statistiques de production."
)

# ── 4. Gestion des récolteurs ─────────────────────────────────────────────────
heading1(doc, "4. Module Gestion des Récolteurs")

body(
    doc,
    "Un récolteur est un travailleur qui effectue la récolte des régimes de palme. "
    "Ce module centralise toutes les informations relatives aux récolteurs."
)

feature_block(
    doc, "👷", "Fiche récolteur détaillée",
    "Chaque récolteur est enregistré avec ses informations complètes :",
    [
        "Nom complet",
        "Lieu de résidence",
        "Numéro de téléphone (identifiant unique)",
        "Numéro WhatsApp et statut d'activation WhatsApp",
        "Inscription Wave (paiement mobile)",
        "Mobile Money : opérateur (Orange Money, MTN MoMo, Moov Money)",
    ]
)

feature_block(
    doc, "📈", "Performance individuelle",
    "Chaque récolteur dispose d'une page de performance qui affiche :",
    [
        "Total de régimes récoltés par période",
        "Évolution mensuelle de sa production",
        "Classement par rapport aux autres récolteurs",
        "Historique de toutes ses participations aux fiches de récolte",
    ]
)

feature_block(
    doc, "📋", "Liste avec filtres et recherche",
    "L'interface liste tous les récolteurs avec une barre de recherche et des filtres "
    "pour retrouver rapidement un récolteur."
)

# ── 5. Fiches de récolte ─────────────────────────────────────────────────────
heading1(doc, "5. Module Fiches de Récolte")

body(
    doc,
    "C'est le cœur de l'application. Une fiche de récolte est le document numérique "
    "qui enregistre toutes les données d'une journée ou d'une période de récolte."
)

feature_block(
    doc, "📝", "Structure complète de la fiche",
    "Chaque fiche de récolte contient :",
    [
        "Date de la récolte",
        "Superviseur général responsable",
        "Liste des superviseurs adjoints avec leurs secteurs",
        "Barème éditable pour chaque type de régime : Grands (FCFA/régime), Moyens, Petits",
        "Dépenses de la journée : nourriture et transport",
        "Observations et remarques libres",
    ]
)

feature_block(
    doc, "📊", "Lignes de récolte par récolteur et par régime",
    "Pour chaque récolteur présent, la fiche enregistre :",
    [
        "Le nom du récolteur (lié à la base ou saisi librement)",
        "Le type de régime récolté : Grands, Moyens ou Petits",
        "La quantité récoltée dans chaque secteur (colonne par secteur)",
    ]
)

feature_block(
    doc, "🧾", "Reçus de vente",
    "Chaque fiche peut contenir un ou plusieurs reçus de vente enregistrant :",
    [
        "Date de la vente",
        "Nom du client acheteur",
        "Poids total pesé en kg",
        "Pourcentage de non-conformes",
        "Montant total en FCFA",
    ]
)

feature_block(
    doc, "✅", "Workflow de validation à 3 états",
    "Chaque fiche passe par un cycle de validation :",
    [
        "Brouillon : en cours de saisie, modifiable",
        "Soumis : transmis au responsable pour validation",
        "Validé : approuvé définitivement, non modifiable",
    ]
)

feature_block(
    doc, "📜", "Historique des récoltes",
    "Une page liste toutes les fiches de récolte avec filtres par date, statut et secteur."
)

# ── 6. Fiches de travaux ─────────────────────────────────────────────────────
heading1(doc, "6. Module Fiches de Travaux")

body(
    doc,
    "Ce module permet d'enregistrer tous les travaux agricoles réalisés sur la palmeraie "
    "(entretien, traitement, fertilisation, etc.)."
)

feature_block(
    doc, "🌿", "Types de travaux pris en charge",
    "L'application couvre 10 natures de travaux différentes :",
    [
        "Désherbage",
        "Traitement phytosanitaire",
        "Fertilisation",
        "Taille / Ablation",
        "Récolte",
        "Pépinière",
        "Plantation",
        "Entretien voie d'accès",
        "Entretien infrastructure",
        "Recensement",
        "Autre (champ libre)",
    ]
)

feature_block(
    doc, "📝", "Contenu d'une fiche de travaux",
    "Chaque fiche de travaux enregistre :",
    [
        "Superviseur responsable des travaux",
        "Nature des travaux réalisés",
        "Superficie couverte en hectares",
        "Période des travaux (dates ou description)",
        "Nombre de personnes mobilisées",
        "Secteurs concernés (liste des secteurs couverts)",
        "Observations et remarques libres",
    ]
)

feature_block(
    doc, "🧪", "Consommables utilisés",
    "Pour chaque fiche, on enregistre les consommables nécessaires :",
    [
        "Désignation du produit ou matériel",
        "Quantité utilisée avec unité",
        "Prix unitaire — calcul automatique du coût total",
    ]
)

feature_block(
    doc, "👷", "Répartition des tâches et main d'œuvre",
    "La fiche détaille la contribution de chaque personne :",
    [
        "Nom et prénom du travailleur",
        "Nature des tâches effectuées",
        "Quantité réalisée et prix unitaire — calcul automatique de la rémunération",
    ]
)

feature_block(
    doc, "✅", "Workflow de validation identique aux récoltes",
    "Même cycle Brouillon → Soumis → Validé que pour les fiches de récolte."
)

# ── 7. Matériels ─────────────────────────────────────────────────────────────
heading1(doc, "7. Module Gestion des Matériels et Équipements")

feature_block(
    doc, "🔧", "Inventaire du matériel",
    "L'application dispose d'un module pour gérer les équipements et matériels de la palmeraie :",
    [
        "Numéro d'inventaire unique",
        "Désignation de l'équipement",
        "Quantité disponible",
        "Suivi de l'état du matériel",
    ]
)

# ── 8. Tableau de bord ───────────────────────────────────────────────────────
heading1(doc, "8. Tableau de Bord (Dashboard)")

body(
    doc,
    "Le tableau de bord est la page d'accueil principale. Il offre une vue synthétique "
    "et graphique de toute l'activité de la palmeraie, avec des filtres dynamiques."
)

feature_block(
    doc, "📊", "Indicateurs clés (KPIs)",
    "Les métriques suivantes sont calculées en temps réel et actualisées selon l'année sélectionnée :",
    [
        "Production totale en régimes",
        "Nombre de secteurs actifs et total",
        "Nombre de récolteurs actifs",
        "Rendement moyen en régimes/hectare",
        "Poids total pesé en kg et poids moyen par régime",
        "Montant total des ventes en FCFA",
        "Total des dépenses de récolte (nourriture + transport)",
        "Coût total des travaux agricoles",
        "Nombre de fiches de récolte et de travaux créées",
    ]
)

feature_block(
    doc, "📈", "Graphiques de production",
    "Plusieurs visualisations graphiques sont disponibles :",
    [
        "Production mensuelle sur les 6 derniers mois glissants",
        "Production mensuelle sur l'année sélectionnée",
        "Comparaison production N vs N-1 (mois par mois)",
        "Production multi-années (courbes superposées)",
        "Production journalière (calendrier de l'année)",
        "Production par secteur (histogramme)",
        "Rendement par secteur (régimes/ha)",
    ]
)

feature_block(
    doc, "💰", "Graphiques financiers",
    "",
    [
        "Montant des ventes mensuel (6 mois et annuel)",
        "Comparaison ventes N vs N-1",
        "Dépenses vs production (mensuel)",
        "Coût des travaux par nature (top 8)",
        "Coût des travaux mensuel (annuel et 6 mois)",
        "Comparaison coûts travaux N vs N-1",
    ]
)

feature_block(
    doc, "🏆", "Performance des récolteurs",
    "",
    [
        "Top 5 des récolteurs les plus productifs (annuel et 6 mois)",
        "Comparaison performance N vs N-1 pour chaque récolteur",
    ]
)

feature_block(
    doc, "🔍", "Filtres dynamiques",
    "Toutes les statistiques du tableau de bord peuvent être filtrées dynamiquement par :",
    [
        "Année (menu déroulant)",
        "Secteur spécifique",
        "Récolteur spécifique",
        "Type de régime (Grands / Moyens / Petits)",
    ]
)

# ── 9. Notifications ─────────────────────────────────────────────────────────
heading1(doc, "9. Module Notifications")

feature_block(
    doc, "🔔", "Centre de notifications",
    "Un système de notifications intégré informe les utilisateurs des événements importants :",
    [
        "Cloche de notification dans la barre de navigation avec compteur",
        "Notifications de type : succès, avertissement, information",
        "Marquage lu / non lu",
        "Lien direct vers l'élément concerné",
    ]
)

# ── 10. Interface générale ────────────────────────────────────────────────────
heading1(doc, "10. Interface Utilisateur et Ergonomie")

feature_block(
    doc, "🖥️", "Design moderne et responsive",
    "L'interface est développée avec React et une bibliothèque de composants moderne :",
    [
        "Sidebar de navigation latérale avec menu structuré",
        "Thème clair avec palette cohérente",
        "Tableaux de données avec tri, filtres et pagination",
        "Dialogues de confirmation avant suppression",
        "Messages de succès et d'erreur contextuels (toasts)",
        "Loader animé pendant le chargement des données",
        "Sélecteurs avec recherche intégrée (SearchableSelect)",
    ]
)

feature_block(
    doc, "🛡️", "Sécurité et contrôle d'accès",
    "",
    [
        "Routes protégées : accès refusé si non connecté",
        "Routes admin : accès refusé si rôle insuffisant",
        "Token JWT stocké et envoyé automatiquement à chaque requête",
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  PARTIE 2 — QUESTIONS AU DIRECTEUR DE MÉMOIRE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PARTIE 2")
set_run_font(r, size=10, bold=True, color=(127, 127, 127))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("QUESTIONS POUR LE DIRECTEUR DE MÉMOIRE")
set_run_font(r, size=16, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Ces questions ont pour objectif de préciser le cahier des charges, "
    "de mieux cerner les attentes et de s'assurer que le développement "
    "répond aux besoins réels du terrain."
)
set_run_font(r, size=11, italic=True, color=(89, 89, 89))

doc.add_paragraph()

# ── Thème A ──────────────────────────────────────────────────────────────────
heading1(doc, "A. Contexte et périmètre de la palmeraie")

question_item(doc, 1,
    "Combien de secteurs compte exactement la palmeraie ?",
    "Cela détermine si la structure actuelle (code + nom + superficie) est suffisante "
    "ou s'il faut ajouter d'autres attributs (zone, bloc, rangée…)."
)
question_item(doc, 2,
    "Combien de récolteurs travaillent sur la palmeraie ? Sont-ils permanents ou saisonniers ?",
    "Cela influence la gestion des fiches (récolteurs ponctuels vs. base stable)."
)
question_item(doc, 3,
    "La palmeraie appartient-elle à une seule entité ou à plusieurs propriétaires ?",
    "En cas de copropriété, il faudrait peut-être gérer des accès séparés par propriétaire."
)
question_item(doc, 4,
    "Y a-t-il plusieurs palmeraies à gérer, ou une seule ?",
    "Si plusieurs sites, l'application doit permettre de les distinguer."
)

# ── Thème B ──────────────────────────────────────────────────────────────────
heading1(doc, "B. Processus métier actuel")

question_item(doc, 5,
    "Comment la gestion se fait-elle actuellement : sur papier, sur Excel, dans un registre ?",
    "Voir les documents existants permettrait d'aligner exactement les champs de l'application."
)
question_item(doc, 6,
    "Qui saisit les données aujourd'hui : le superviseur, un agent de terrain, un secrétaire ?",
    "Cela précise qui sont les vrais utilisateurs quotidiens."
)
question_item(doc, 7,
    "Quelles sont les erreurs ou pertes d'information les plus fréquentes dans le processus actuel ?",
    "Ces points de douleur sont exactement ce que l'application doit résoudre en priorité."
)
question_item(doc, 8,
    "Existe-t-il des fiches papier physiques servant de référence ? Si oui, puis-je les consulter ?",
    "Les fiches papier existantes doivent être reflétées fidèlement dans l'application numérique."
)

# ── Thème C ──────────────────────────────────────────────────────────────────
heading1(doc, "C. Utilisateurs et rôles")

question_item(doc, 9,
    "Qui va utiliser cette application au quotidien ? Combien de personnes, quels profils ?",
    "Actuellement deux rôles sont implémentés : Administrateur et Superviseur."
)
question_item(doc, 10,
    "Les récolteurs eux-mêmes ont-ils besoin d'accéder à l'application ?",
    "Si oui, il faudrait un troisième rôle avec une interface simplifiée."
)
question_item(doc, 11,
    "Y a-t-il un directeur ou propriétaire qui doit consulter des rapports sans modifier les données ?",
    "Ce profil 'lecture seule' n'est pas encore implémenté."
)
question_item(doc, 12,
    "Le directeur de mémoire lui-même utilisera-t-il l'application, ou est-il uniquement commanditaire académique ?",
    "Cela précise si une démonstration sur données réelles est attendue."
)

# ── Thème D ──────────────────────────────────────────────────────────────────
heading1(doc, "D. Fonctionnalités prioritaires et manquantes")

question_item(doc, 13,
    "Parmi les modules déjà développés (voir Partie 1), lesquels sont les plus importants pour vous ?",
    "Cela permet de concentrer les efforts de finalisation sur l'essentiel."
)
question_item(doc, 14,
    "Y a-t-il une fonctionnalité que vous attendiez et que vous ne voyez pas dans ce qui a été présenté ?",
    "Question ouverte pour identifier les fonctionnalités manquantes."
)
question_item(doc, 15,
    "La gestion des paiements des récolteurs doit-elle être intégrée dans l'application ?",
    "Le modèle récolteur contient déjà les informations de paiement mobile (Wave, MTN MoMo…). "
    "Faut-il aller plus loin avec un module paiement complet ?"
)
question_item(doc, 16,
    "Les alertes automatiques (stock matériel bas, travaux en retard, objectif de production non atteint) sont-elles importantes ?",
    "Cela orienterait le développement du module notifications."
)

# ── Thème E ──────────────────────────────────────────────────────────────────
heading1(doc, "E. Données, rapports et exports")

question_item(doc, 17,
    "Quels rapports ou états de sortie attendez-vous de l'application ?",
    "Exemples : rapport mensuel de production, état des paiements, bilan annuel en PDF ou Excel."
)
question_item(doc, 18,
    "Quels indicateurs clés (KPIs) souhaitez-vous absolument voir sur le tableau de bord ?",
    "Les KPIs actuels sont listés en Partie 1. Y en a-t-il d'autres à ajouter ?"
)
question_item(doc, 19,
    "Faut-il pouvoir exporter les données vers Excel pour une analyse externe ?",
    "Cette fonctionnalité n'est pas encore implémentée."
)
question_item(doc, 20,
    "Les comparaisons annuelles (année N vs N-1) déjà implémentées vous sont-elles utiles ?",
    "Si non, elles peuvent être remplacées par d'autres visualisations plus pertinentes."
)

# ── Thème F ──────────────────────────────────────────────────────────────────
heading1(doc, "F. Contraintes techniques et déploiement")

question_item(doc, 21,
    "L'application sera-t-elle utilisée en ligne (internet) ou hors ligne (réseau local) ?",
    "Une utilisation sur le terrain sans connexion nécessiterait une architecture différente."
)
question_item(doc, 22,
    "Les utilisateurs auront-ils des smartphones, tablettes ou uniquement des ordinateurs ?",
    "L'interface est déjà responsive, mais une version mobile dédiée peut être nécessaire."
)
question_item(doc, 23,
    "Y a-t-il des exigences sur la langue de l'interface ?",
    "L'application est actuellement entièrement en français."
)
question_item(doc, 24,
    "Faut-il prévoir une connexion avec un système existant (logiciel comptable, ERP, autre base de données) ?",
    "Si oui, des API d'intégration seraient à développer."
)

# ── Thème G ──────────────────────────────────────────────────────────────────
heading1(doc, "G. Ce qui ne vous convainc pas (retour direct)")

question_item(doc, 25,
    "Quand vous regardez ce qui a été développé, qu'est-ce qui concrètement ne vous convient pas ?",
    "Ergonomie ? Données enregistrées ? Fonctionnalités manquantes ? Technologie utilisée ?"
)
question_item(doc, 26,
    "Si vous deviez retirer 3 choses et ajouter 3 choses, ce serait lesquelles ?",
    "Cette question force à prioriser clairement."
)
question_item(doc, 27,
    "Y a-t-il une application existante (même dans un autre domaine) dont vous appréciez le fonctionnement et qui pourrait servir de modèle ?",
    "Une référence concrète vaut mieux que des descriptions abstraites."
)

# ── Note finale ───────────────────────────────────────────────────────────────
doc.add_paragraph()
add_separator(doc)
doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("Note de méthode")
set_run_font(r, size=11, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
r = p.add_run(
    "Avant la réunion, il est recommandé de préparer une démonstration courte (5 à 10 minutes) "
    "des fonctionnalités principales : tableau de bord, saisie d'une fiche de récolte, "
    "liste des récolteurs et des secteurs. Une démonstration sur des données réelles "
    "ou représentatives rendra les échanges beaucoup plus concrets et permettra "
    "au directeur de réagir sur du réel plutôt que sur des descriptions abstraites."
)
set_run_font(r, size=11, italic=True)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
output_path = r"e:\UFHB\MASTER2_BDGL\Projet_Memoire\theme2\Projet_Gestion_Palmeraie\Document_Projet_Palmeraie.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
